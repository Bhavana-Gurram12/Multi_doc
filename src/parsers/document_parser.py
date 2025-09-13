import PyPDF2
import docx
from bs4 import BeautifulSoup
from typing import Dict, Tuple

try:
    import magic
    HAS_MAGIC = True
except ImportError:
    HAS_MAGIC = False

class DocumentParser:
    """Multi-format document parser"""
    
    def detect_format(self, file_path: str) -> str:
        """Detect document format"""
        if HAS_MAGIC:
            try:
                mime = magic.from_file(file_path, mime=True)
                if 'pdf' in mime:
                    return 'pdf'
                elif 'word' in mime or file_path.endswith('.docx'):
                    return 'docx'
                elif 'html' in mime or file_path.endswith('.html'):
                    return 'html'
                elif 'text' in mime:
                    return 'text'
            except:
                pass
        
        # Fallback to extension
        ext = file_path.split('.')[-1].lower()
        return ext if ext in ['pdf', 'docx', 'html', 'txt'] else 'unknown'
    
    def parse_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """Parse PDF document"""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            content = ""
            for page in reader.pages:
                content += page.extract_text() + "\n"
            
            metadata = {
                'pages': len(reader.pages),
                'title': reader.metadata.get('/Title', '') if reader.metadata else ''
            }
            return content.strip(), metadata
    
    def parse_docx(self, file_path: str) -> Tuple[str, Dict]:
        """Parse DOCX document"""
        doc = docx.Document(file_path)
        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        metadata = {
            'paragraphs': len(doc.paragraphs),
            'title': doc.core_properties.title or ''
        }
        return content.strip(), metadata
    
    def parse_html(self, file_path: str) -> Tuple[str, Dict]:
        """Parse HTML document"""
        with open(file_path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file.read(), 'html.parser')
            content = soup.get_text()
            
            metadata = {
                'title': soup.title.string if soup.title else '',
                'links': len(soup.find_all('a'))
            }
            return content.strip(), metadata
    
    def parse_text(self, file_path: str) -> Tuple[str, Dict]:
        """Parse plain text document"""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            
            metadata = {
                'lines': len(content.split('\n')),
                'chars': len(content)
            }
            return content.strip(), metadata
    
    def parse_document(self, file_path: str) -> Tuple[str, Dict, str]:
        """Parse any supported document format"""
        doc_type = self.detect_format(file_path)
        
        if doc_type == 'pdf':
            content, metadata = self.parse_pdf(file_path)
        elif doc_type == 'docx':
            content, metadata = self.parse_docx(file_path)
        elif doc_type == 'html':
            content, metadata = self.parse_html(file_path)
        elif doc_type in ['text', 'txt']:
            content, metadata = self.parse_text(file_path)
        else:
            content, metadata = "", {}
        
        return content, metadata, doc_type